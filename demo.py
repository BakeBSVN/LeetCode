import fitz  # PyMuPDF
from googletrans import Translator
import docx as aw

# Function to translate text to Vietnamese using Google Translate
def translate_to_vietnamese(text):
    translator = Translator()
    translated_text = translator.translate(text, src='en', dest='vi').text
    return translated_text

# Open the PDF file
with fitz.open("Input.pdf") as doc:
    # Get the text from the first page
    text = doc.load_page(0).get_text()

    # Translate the text to Vietnamese
    vietnamese_text = translate_to_vietnamese(text)

    # Create a new Word document with the translated text
    doc = aw.Document()
    page = doc.add_paragraph()
    run = page.add_run(vietnamese_text)
    run.font.name = "Times New Roman"
    run.font.size = aw.shared.Pt(12)

    # Save the Word document
    doc.save("Output2.docx")
